#!/usr/bin/env python3
"""
Convert AVATAR OnE 플랫폼 기능명세서 v1.5 markdown to DOCX.

Replicates the exact layout/design system of the v1.2 docx file:
- Cover page with centered title, subtitle, metadata table
- Section cards (numbered blue/light-blue 1x2 tables)
- Feature tables with dark navy headers
- Callout boxes (yellow/blue backgrounds)
- JSON code block in light gray box
- Footer page numbers
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Color palette (v1.2 design system)
# ---------------------------------------------------------------------------
CLR_DARK_NAVY = "1A3C5E"
CLR_MED_BLUE = "2874A6"
CLR_DARK_BLUEGRAY = "2C3E50"
CLR_RED = "C0392B"
CLR_ORANGE = "D4740A"
CLR_GREEN = "27AE60"
CLR_WHITE = "FFFFFF"
CLR_LIGHT_BLUE = "EAF2F8"
CLR_LIGHT_YELLOW = "FEF5E7"
CLR_VERY_LIGHT_GRAY = "F4F6F7"

# RGB helpers
RGB_DARK_NAVY = RGBColor(0x1A, 0x3C, 0x5E)
RGB_MED_BLUE = RGBColor(0x28, 0x74, 0xA6)
RGB_DARK_BLUEGRAY = RGBColor(0x2C, 0x3E, 0x50)
RGB_RED = RGBColor(0xC0, 0x39, 0x2B)
RGB_ORANGE = RGBColor(0xD4, 0x74, 0x0A)
RGB_GREEN = RGBColor(0x27, 0xAE, 0x60)
RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Font
FONT_MAIN = "Arial Unicode MS"
FONT_CODE = "Courier New"


# ===== Low-level helpers ====================================================

def _set_cell_shading(cell, color_hex):
    """Set cell background fill colour."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_width(cell, width_twips):
    """Set preferred cell width in twips."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_cell_vertical_alignment(cell, val="center"):
    tc_pr = cell._element.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), val)
    tc_pr.append(v_align)


def _set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _set_paragraph_spacing(para, before=None, after=None, line=None):
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))


def _make_run(para, text, font_name=FONT_MAIN, size_pt=None, bold=None,
              color_rgb=None, italic=None):
    """Add a run with explicit formatting."""
    run = para.add_run(text)
    run.font.name = font_name
    # East-Asian font fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    if italic is not None:
        run.italic = italic
    return run


def _set_table_borders(table, sz=4, color="000000"):
    """Set single borders on all sides of a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _set_table_width(table, width_twips):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    # Remove existing tblW if any
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.insert(0, tbl_w)


def _set_table_alignment(table, alignment="center"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), alignment)
    existing = tbl_pr.find(qn("w:jc"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(jc)


def _set_table_grid(table, col_widths_twips):
    """Set tblGrid with explicit column widths."""
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tbl.insert(1, grid)


def _add_page_number_footer(section):
    """Add centered page number field to section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    run.font.size = Pt(9)
    run.font.name = FONT_MAIN


# ===== Markdown parsing =====================================================

def _parse_md_table(lines):
    """Parse markdown table lines into list of row-lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith("|---") or re.match(r"^\|[\s\-:]+\|$", line):
            continue
        # Handle escaped pipes
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    return rows


def _parse_markdown(md_text):
    """
    Parse the v1.3 markdown into a structured list of blocks.
    Returns: list of dicts with type + data keys.
    """
    lines = md_text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Code block ---
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            i += 1
            continue

        # --- Table ---
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_md_table(table_lines)
            blocks.append({"type": "table", "rows": rows})
            continue

        # --- Headings ---
        if stripped.startswith("#### "):
            blocks.append({"type": "h4", "text": stripped[5:]})
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:]})
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:]})
            i += 1
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:]})
            i += 1
            continue

        # --- HR ---
        if stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        # --- List item ---
        if stripped.startswith("- "):
            blocks.append({"type": "list_item", "text": stripped[2:]})
            i += 1
            continue

        # --- Paragraph text ---
        if stripped:
            # Collect consecutive non-empty lines into a paragraph
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                s = lines[i].strip()
                if (not s or s.startswith("#") or s.startswith("|") or
                        s.startswith("```") or s.startswith("- ") or s == "---"):
                    break
                para_lines.append(s)
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue

        # Empty line
        i += 1

    return blocks


# ===== Document construction ================================================

def _add_cover_page(doc):
    """Build the cover page (section 0)."""
    section = doc.sections[0]
    section.top_margin = Twips(1440)     # 1 inch
    section.bottom_margin = Twips(1440)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)

    # P[0]: Top spacer
    p0 = doc.add_paragraph()
    _set_paragraph_spacing(p0, before=2800)

    # P[1]: Main title
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(p1, "AVATAR OnE \ud50c\ub7ab\ud3fc", size_pt=28, bold=True, color_rgb=RGB_DARK_NAVY)
    _set_paragraph_spacing(p1, after=120)

    # P[2]: Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(p2, "\uae30\ub2a5 \uba85\uc138\uc11c", size_pt=18, bold=False, color_rgb=RGB_MED_BLUE)
    _set_paragraph_spacing(p2, after=500)

    # P[3]: Spacer
    p3 = doc.add_paragraph()
    _set_paragraph_spacing(p3, before=80, after=80)

    # P[4]: Spacer
    p4 = doc.add_paragraph()
    _set_paragraph_spacing(p4, before=200)

    # P[5]: Workflow line
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(p5, "Builder \u2192 Trainer \u2192 \ud14c\uc2a4\ud2b8/\uc2b9\uc778 \u2192 \uc2a4\ucf00\uc904\ub9c1 \u2192 \uacb0\uacfc \uc870\ud68c",
              size_pt=11, color_rgb=RGB_DARK_BLUEGRAY)
    _set_paragraph_spacing(p5, after=80)

    # P[6]: Sub-description
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(p6, "\uc804\uccb4 \uc6cc\ud06c\ud50c\ub85c\uc6b0 \uae30\ubc18 \uae30\ub2a5 \uba85\uc138", size_pt=11,
              color_rgb=RGB_DARK_BLUEGRAY)
    _set_paragraph_spacing(p6, after=400)

    # P[7]: Spacer
    p7 = doc.add_paragraph()
    _set_paragraph_spacing(p7, before=400)

    # Metadata table (4 x 2)
    meta_data = [
        ("\ubc84\uc804", "1.5"),
        ("\uc791\uc131\uc77c", "2025-02-02"),
        ("\uc218\uc815\uc77c", "2026-02-09"),
        ("\ub300\uc0c1", "\uace0\uac1d \uc804\ub2ec\uc6a9"),
    ]
    tbl = doc.add_table(rows=len(meta_data), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_grid(tbl, [1600, 3000])
    _set_table_width(tbl, 4600)

    for idx, (label, value) in enumerate(meta_data):
        left = tbl.rows[idx].cells[0]
        right = tbl.rows[idx].cells[1]

        _set_cell_width(left, 1600)
        _set_cell_width(right, 3000)

        _set_cell_shading(left, CLR_LIGHT_BLUE)
        _set_cell_vertical_alignment(left, "center")
        _set_cell_vertical_alignment(right, "center")

        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _make_run(lp, label, size_pt=9, bold=True, color_rgb=RGB_DARK_BLUEGRAY)

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _make_run(rp, value, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)


def _add_section_card(doc, number, title, description):
    """Add a section card table: 1x2 with number on the left, title/desc on the right."""
    # Add Heading 1 (invisible bookmark target for TOC)
    h1 = doc.add_heading(title, level=1)
    for run in h1.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGB_DARK_NAVY
        run.font.name = FONT_MAIN
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    # Make the heading invisible (font size 1) to avoid duplicate display
    for run in h1.runs:
        run.font.size = Pt(1)
        run.font.color.rgb = RGB_WHITE

    tbl = doc.add_table(rows=1, cols=2)
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_grid(tbl, [700, 8660])
    _set_table_width(tbl, 9360)

    # Left cell: section number
    left = tbl.rows[0].cells[0]
    _set_cell_shading(left, CLR_MED_BLUE)
    _set_cell_width(left, 700)
    _set_cell_vertical_alignment(left, "center")
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(lp, str(number), size_pt=11, bold=True, color_rgb=RGB_WHITE)

    # Right cell: title + description
    right = tbl.rows[0].cells[1]
    _set_cell_shading(right, CLR_LIGHT_BLUE)
    _set_cell_width(right, 8660)
    _set_cell_vertical_alignment(right, "center")
    _set_cell_margins(right, top=60, bottom=60, left=120, right=80)

    rp0 = right.paragraphs[0]
    _make_run(rp0, title, size_pt=11, bold=True, color_rgb=RGB_DARK_BLUEGRAY)

    rp1 = right.add_paragraph()
    _make_run(rp1, description, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)

    # Small spacing after card
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=60, after=60)


def _add_callout_box(doc, title, body_text, fill_color):
    """Add a 1x1 table callout box with coloured background."""
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_width(tbl, 9360)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, fill_color)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    if title:
        tp = cell.paragraphs[0]
        _make_run(tp, title, size_pt=10, bold=True, color_rgb=RGB_DARK_BLUEGRAY)
        _set_paragraph_spacing(tp, after=60)

        bp = cell.add_paragraph()
    else:
        bp = cell.paragraphs[0]

    # Handle multi-paragraph body: split on double-newline or \n\n
    parts = body_text.split("\n\n") if "\n\n" in body_text else [body_text]
    first = True
    for part in parts:
        if first:
            _add_rich_text(bp, part.strip(), size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)
            first = False
        else:
            np = cell.add_paragraph()
            _add_rich_text(np, part.strip(), size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)

    # Spacer after box
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=40, after=40)


def _add_rich_text(para, text, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY):
    """Add text with **bold** markdown support."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _make_run(para, part[2:-2], size_pt=size_pt, bold=True, color_rgb=color_rgb)
        else:
            _make_run(para, part, size_pt=size_pt, color_rgb=color_rgb)


def _add_feature_table(doc, rows):
    """
    Add a feature table (ID/기능명/설명/우선순위/구현단계).
    rows[0] = header, rows[1:] = data.
    """
    col_widths = [700, 2000, 4360, 700, 1600]
    num_cols = len(col_widths)

    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_grid(tbl, col_widths)
    _set_table_width(tbl, sum(col_widths))

    # Header row
    for ci in range(num_cols):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, CLR_DARK_NAVY)
        _set_cell_width(cell, col_widths[ci])
        _set_cell_vertical_alignment(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = rows[0][ci] if ci < len(rows[0]) else ""
        _make_run(p, text, size_pt=9, bold=True, color_rgb=RGB_WHITE)

    # Data rows
    for ri in range(1, len(rows)):
        for ci in range(num_cols):
            cell = tbl.rows[ri].cells[ci]
            _set_cell_width(cell, col_widths[ci])
            _set_cell_vertical_alignment(cell, "center")
            text = rows[ri][ci] if ci < len(rows[ri]) else ""
            p = cell.paragraphs[0]

            if ci == 0:  # ID
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _make_run(p, text, size_pt=9, bold=True, color_rgb=RGB_DARK_BLUEGRAY)
            elif ci == 3:  # Priority
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                clr = RGB_RED if text.strip() == "P0" else RGB_ORANGE if text.strip() == "P1" else RGB_DARK_BLUEGRAY
                _make_run(p, text, size_pt=9, bold=True, color_rgb=clr)
            elif ci == 4:  # Phase
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _make_run(p, text, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)
            else:  # Name, Description
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _make_run(p, text, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)

    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=40, after=40)


def _add_generic_table(doc, rows, col_widths=None):
    """
    Add a non-feature table with dark-navy header.
    Handles special formatting for priority-like values.
    """
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    # Pad rows
    for r in rows:
        while len(r) < num_cols:
            r.append("")

    if col_widths is None:
        total = 9360
        col_widths = [total // num_cols] * num_cols
        # Adjust last column to absorb remainder
        col_widths[-1] = total - sum(col_widths[:-1])

    # Ensure col_widths matches num_cols
    if len(col_widths) < num_cols:
        extra = num_cols - len(col_widths)
        col_widths.extend([col_widths[-1]] * extra)
    elif len(col_widths) > num_cols:
        col_widths = col_widths[:num_cols]

    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_grid(tbl, col_widths)
    _set_table_width(tbl, sum(col_widths))

    # Header row
    for ci in range(num_cols):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, CLR_DARK_NAVY)
        _set_cell_width(cell, col_widths[ci])
        _set_cell_vertical_alignment(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _make_run(p, rows[0][ci], size_pt=9, bold=True, color_rgb=RGB_WHITE)

    # Data rows
    for ri in range(1, len(rows)):
        for ci in range(num_cols):
            cell = tbl.rows[ri].cells[ci]
            _set_cell_width(cell, col_widths[ci])
            _set_cell_vertical_alignment(cell, "center")
            text = rows[ri][ci]
            p = cell.paragraphs[0]

            # First column bold
            is_bold = (ci == 0)

            # Special colour for priority values
            clr = RGB_DARK_BLUEGRAY
            t = text.strip()
            if t == "\ub192\uc74c":
                clr = RGB_RED
                is_bold = True
            elif t == "\ub0ae\uc74c":
                clr = RGB_GREEN
                is_bold = True

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _make_run(p, text, size_pt=9, bold=is_bold, color_rgb=clr)

    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=40, after=40)


def _add_json_code_block(doc, code_text):
    """Add JSON code as a light-gray 1x1 table box."""
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders(tbl, sz=4, color="000000")
    _set_table_width(tbl, 9360)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, CLR_VERY_LIGHT_GRAY)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    p = cell.paragraphs[0]
    _make_run(p, code_text, font_name=FONT_CODE, size_pt=9, color_rgb=RGB_DARK_BLUEGRAY)
    _set_paragraph_spacing(p, before=0, after=0)

    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=40, after=40)


def _add_heading2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGB_DARK_BLUEGRAY
        run.font.name = FONT_MAIN
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    return h


def _add_normal_paragraph(doc, text, size_pt=9):
    p = doc.add_paragraph()
    _add_rich_text(p, text, size_pt=size_pt, color_rgb=RGB_DARK_BLUEGRAY)
    return p


def _add_list_item(doc, text, size_pt=9):
    p = doc.add_paragraph()
    _add_rich_text(p, "- " + text, size_pt=size_pt, color_rgb=RGB_DARK_BLUEGRAY)
    return p


# ===== Main build logic =====================================================

def _build_document(md_text):
    """Build the complete DOCX document from markdown text."""
    doc = Document()
    blocks = _parse_markdown(md_text)

    # ---- Workflow descriptions (from section 1 table) for section card use ----
    section_descriptions = {}
    # Find the workflow table (it's the first table after ## 1. heading)
    wf_table = None
    for bi, b in enumerate(blocks):
        if b["type"] == "h2" and b["text"].startswith("1."):
            # Look for next table
            for j in range(bi + 1, len(blocks)):
                if blocks[j]["type"] == "table":
                    wf_table = blocks[j]["rows"]
                    break
                if blocks[j]["type"] == "h2":
                    break
            break

    if wf_table and len(wf_table) > 1:
        # Map step number -> description
        for row in wf_table[1:]:
            if len(row) >= 3:
                step_num = row[0].strip()
                desc = row[2].strip()
                section_descriptions[step_num] = desc

    # Section mapping: section_number -> (title, description)
    # Sections 1-6 map to workflow table steps 1-6
    sec_map = {
        1: ("\uc804\uccb4 \uc6cc\ud06c\ud50c\ub85c\uc6b0", "\ubcf8 \ud50c\ub7ab\ud3fc\uc758 \ud575\uc2ec \uc6cc\ud06c\ud50c\ub85c\uc6b0\ub97c \ub2e4\uc74c\uacfc \uac19\uc774 \uc815\uc758\ud55c\ub2e4."),
        2: ("Builder (App \uac1c\ubc1c)", section_descriptions.get("1", "")),
        3: ("\ucef4\ud3ec\ub10c\ud2b8 \uae00\ub85c\ubc8c \ub77c\uc774\ube0c\ub7ec\ub9ac", "\uc2dc\uc2a4\ud15c \uc804\uc5ed\uc5d0\uc11c \uc7ac\uc0ac\uc6a9 \uac00\ub2a5\ud55c \ucef4\ud3ec\ub10c\ud2b8\ub97c \ub4f1\ub85d\u00b7\uad00\ub9ac\ud558\ub294 \uce74\ud0c8\ub85c\uadf8\uc774\ub2e4."),
        4: ("Trainer (\ud559\uc2b5 \uc694\uccad)", "\uac1c\ubc1c \uc644\ub8cc\ub41c App\uc5d0 \ub300\ud574 \ud559\uc2b5 \ud30c\ub77c\ubbf8\ud130\uc640 \ub9ac\uc18c\uc2a4\ub97c \uc124\uc815\ud558\uc5ec \ud559\uc2b5 \uc694\uccad\uc744 \uc81c\ucd9c\ud558\ub294 \uc804\uc6a9 \ud654\uba74\uc774\ub2e4."),
        5: ("\ud14c\uc2a4\ud2b8 \ubc0f \uc2b9\uc778", section_descriptions.get("3", "")),
        6: ("\uc2e4\ud589 \ub300\uae30\uc5f4 \ubc0f \uc6b0\uc120\uc21c\uc704 \uad00\ub9ac", section_descriptions.get("4", "")),
        7: ("\ub9ac\uc18c\uc2a4 \uad00\ub9ac", section_descriptions.get("5", "")),
        8: ("\uacb0\uacfc \ubaa8\ub378 \uad00\ub9ac", section_descriptions.get("6", "")),
        9: ("\uc6cc\ud06c\ub85c\ub4dc \uc2a4\ucf00\uc904\ub9c1 (HPC)", "\uc2b9\uc778\ub41c \uc6cc\ud06c\ub85c\ub4dc\ub97c \uc2e4\uc81c HPC \uc778\ud504\ub77c\uc5d0\uc11c \uc2e4\ud589\ud558\ub294 \uacc4\uce35\uc774\ub2e4."),
        10: ("\uc800\uc7a5\uc18c", "\uc2a4\ud399 \ud30c\uc77c, \ubaa8\ub378, \ud14c\uc2a4\ud2b8 \uc2e4\ud589 \uacb0\uacfc, \uc2e4\ud589 \ub85c\uadf8 \ub4f1\uc744 \uc800\uc7a5\ud558\uace0 \uad00\ub9ac\ud55c\ub2e4."),
        11: ("\uad8c\ud55c \uad00\ub9ac \ubc0f \uc778\uc99d", "\uc5ed\ud560 \uae30\ubc18 \uc811\uadfc \uc81c\uc5b4\uc640 \uc0ac\uc6a9\uc790 \uacc4\uc815 \uad00\ub9ac\ub97c \uc81c\uacf5\ud55c\ub2e4."),
        12: ("\uad8c\uc7a5 \uae30\uc220 \uc2a4\ud0dd \ubc0f \uad6c\ud604 \ub85c\ub4dc\ub9f5", "HPC/ML \uc6cc\ud06c\ub85c\ub4dc \uad00\ub9ac\ub97c \uc704\ud55c \uad8c\uc7a5 \uae30\uc220 \uc2a4\ud0dd\uacfc \uad6c\ud604 \ub85c\ub4dc\ub9f5\uc744 \uc815\uc758\ud55c\ub2e4."),
    }

    # ===== COVER PAGE =====
    _add_cover_page(doc)

    # ===== SECTION BREAK =====
    new_section = doc.add_section()
    new_section.top_margin = Twips(1440)
    new_section.bottom_margin = Twips(1440)
    new_section.left_margin = Twips(1440)
    new_section.right_margin = Twips(1440)
    _add_page_number_footer(new_section)

    # ===== CONTENT PAGES =====
    # We iterate through blocks and build content according to section mapping
    current_section_num = 0
    i = 0

    while i < len(blocks):
        b = blocks[i]

        # Skip the h1 (cover title) -- already rendered on cover
        if b["type"] == "h1":
            i += 1
            continue

        # Skip cover-area content (paragraphs / tables before first ##)
        if b["type"] in ("paragraph", "table", "hr") and current_section_num == 0:
            i += 1
            continue

        # ---- Section heading (##) ----
        if b["type"] == "h2":
            m = re.match(r"(\d+)\.\s*(.*)", b["text"])
            if m:
                current_section_num = int(m.group(1))
                sec_title = m.group(2).strip()
                title_full = f"{current_section_num}. {sec_title}"
                desc = sec_map.get(current_section_num, (sec_title, ""))[1]
                _add_section_card(doc, current_section_num, title_full, desc)
            else:
                _add_heading2(doc, b["text"])
            i += 1
            continue

        # ---- Sub-section heading (###) ----
        if b["type"] == "h3":
            text = b["text"].strip()

            # Callout boxes for specific subsections
            if text == "\ud575\uc2ec \uac1c\ub150":
                # Collect following paragraphs until next heading or table
                body_parts = []
                i += 1
                while i < len(blocks) and blocks[i]["type"] in ("paragraph", "list_item"):
                    if blocks[i]["type"] == "paragraph":
                        body_parts.append(blocks[i]["text"])
                    elif blocks[i]["type"] == "list_item":
                        body_parts.append("- " + blocks[i]["text"])
                    i += 1
                body = "\n\n".join(body_parts)
                _add_callout_box(doc, "\ud575\uc2ec \uac1c\ub150", body, CLR_LIGHT_YELLOW)
                continue

            elif text == "\uc6b4\uc601 \ubc94\uc704":
                body_parts = []
                i += 1
                while i < len(blocks) and blocks[i]["type"] in ("paragraph", "list_item"):
                    if blocks[i]["type"] == "paragraph":
                        body_parts.append(blocks[i]["text"])
                    elif blocks[i]["type"] == "list_item":
                        body_parts.append("- " + blocks[i]["text"])
                    i += 1
                body = "\n\n".join(body_parts)
                _add_callout_box(doc, "\uc6b4\uc601 \ubc94\uc704", body, CLR_LIGHT_BLUE)
                continue

            elif text == "RL \ud559\uc2b5 \ud30c\ub77c\ubbf8\ud130":
                body_parts = []
                i += 1
                while i < len(blocks) and blocks[i]["type"] in ("paragraph", "list_item"):
                    if blocks[i]["type"] == "paragraph":
                        body_parts.append(blocks[i]["text"])
                    elif blocks[i]["type"] == "list_item":
                        body_parts.append("- " + blocks[i]["text"])
                    i += 1
                body = "\n\n".join(body_parts)
                _add_callout_box(doc, "RL \ud559\uc2b5 \ud30c\ub77c\ubbf8\ud130", body, CLR_LIGHT_YELLOW)
                continue

            else:
                _add_heading2(doc, text)
                i += 1
                continue

        # ---- Sub-sub-section heading (####) ----
        if b["type"] == "h4":
            text = b["text"].strip()

            # Callout for "테스트 및 승인 프로세스 개요"
            if "\ud14c\uc2a4\ud2b8 \ubc0f \uc2b9\uc778 \ud504\ub85c\uc138\uc2a4 \uac1c\uc694" in text or "\ud504\ub85c\uc138\uc2a4 \uac1c\uc694" in text:
                body_parts = []
                i += 1
                while i < len(blocks) and blocks[i]["type"] in ("paragraph", "list_item"):
                    if blocks[i]["type"] == "paragraph":
                        body_parts.append(blocks[i]["text"])
                    elif blocks[i]["type"] == "list_item":
                        body_parts.append("- " + blocks[i]["text"])
                    i += 1
                body = "\n\n".join(body_parts)
                _add_callout_box(doc, "\ud14c\uc2a4\ud2b8 \ubc0f \uc2b9\uc778 \ud504\ub85c\uc138\uc2a4 \uac1c\uc694", body, CLR_LIGHT_YELLOW)
                continue
            else:
                _add_heading2(doc, text)
                i += 1
                continue

        # ---- Table ----
        if b["type"] == "table":
            rows = b["rows"]
            if not rows:
                i += 1
                continue

            header = [c.strip() for c in rows[0]]

            # Detect feature table by header content
            is_feature = (
                len(header) >= 5
                and "ID" in header[0]
                and ("\uae30\ub2a5\uba85" in header[1] or "\uae30\ub2a5" in header[1])
            )

            if is_feature:
                _add_feature_table(doc, rows)
            else:
                # Determine col widths based on number of columns
                nc = len(header)
                if nc == 2:
                    cw = [2400, 6960]
                elif nc == 3:
                    if "\ub2e8\uacc4" in header[0]:
                        cw = [600, 1800, 6960]
                    elif "\uacc4\uce35" in header[0]:
                        cw = [2200, 2200, 4960]
                    else:
                        cw = [2400, 3480, 3480]
                elif nc == 5:
                    cw = [1400, 1600, 2160, 2000, 2200]
                elif nc == 6:
                    cw = [1000, 2000, 1400, 1400, 1400, 2160]
                elif nc == 7:
                    cw = [600, 1800, 1100, 1600, 1000, 1060, 2200]
                else:
                    total = 9360
                    cw = [total // nc] * nc
                    cw[-1] = total - sum(cw[:-1])

                _add_generic_table(doc, rows, col_widths=cw)
            i += 1
            continue

        # ---- Code block ----
        if b["type"] == "code":
            _add_json_code_block(doc, b["text"])
            i += 1
            continue

        # ---- Paragraph ----
        if b["type"] == "paragraph":
            _add_normal_paragraph(doc, b["text"])
            i += 1
            continue

        # ---- List item ----
        if b["type"] == "list_item":
            _add_list_item(doc, b["text"])
            i += 1
            continue

        # ---- HR ----
        if b["type"] == "hr":
            # Skip -- HRs are just section separators in the markdown
            i += 1
            continue

        i += 1

    return doc


# ===== Entry point ==========================================================

def main():
    base_dir = Path(__file__).parent.parent
    md_file = base_dir / "docs/AVATAR_OnE_\ud50c\ub7ab\ud3fc_\uae30\ub2a5\uba85\uc138\uc11c_v1_5.md"
    docx_file = base_dir / "docs/AVATAR_OnE_\ud50c\ub7ab\ud3fc_\uae30\ub2a5\uba85\uc138\uc11c_v1_5.docx"

    if not md_file.exists():
        print(f"Error: markdown file not found: {md_file}")
        return 1

    md_text = md_file.read_text(encoding="utf-8")
    doc = _build_document(md_text)
    doc.save(str(docx_file))

    if docx_file.exists():
        size_kb = docx_file.stat().st_size / 1024
        print(f"Created: {docx_file}")
        print(f"Size: {size_kb:.1f} KB")
        return 0
    else:
        print("Failed to create output file")
        return 1


if __name__ == "__main__":
    exit(main())
